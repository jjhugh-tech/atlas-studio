from __future__ import annotations

from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSET_DIR = ROOT / "work" / "user-guide" / "assets"
OUTPUT = ROOT / "outputs" / "Atlas_Studio_Complete_User_Guide.docx"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "0A1830"
DEEP = "07111E"
BLUE = "2E74B5"
CYAN = "20B9D5"
PURPLE = "6957D9"
INK = "18263A"
MUTED = "61738B"
PALE = "E8EEF5"
PALE_CYAN = "E8F8FB"
PALE_GOLD = "FFF4D9"
PALE_RED = "FCEBEA"
GREEN = "16845B"
RED = "A23B3B"
WHITE = "FFFFFF"
LINE = "C8D5E5"


AGENTS = [
    ("Atlas", "Platform Intelligence Orchestrator", "Read-only", "diagnostics, research, investigation, memory read, files read", "Receives direction, maintains platform awareness, coordinates approved work."),
    ("Forge", "Platform Development AI", "Approval required", "memory read, files read/write, code execute, test execute", "Primary implementation assistant inside an approved isolated workspace."),
    ("Sentinel", "Security Engineering", "Read-only; approval required", "diagnostics, investigation, memory/files read, security scan", "Threat modeling, secure-code review, dependency analysis, vulnerability triage."),
    ("Verity", "GRC and Compliance", "Read-only", "research, investigation, memory/files read, compliance review, documents", "Control mapping, governance, risk, compliance evidence."),
    ("Quanta", "Quality and Test Engineering", "Approval required", "diagnostics, memory/files read/write, code/test execute", "Test plans, authorized automated tests, release validation, regression gates."),
    ("Sage", "Research and Development", "Read-only", "research, investigation, memory/files read, browser", "Approved research, technology evaluation, experiments, recommendations."),
    ("Counsel", "AI Legal Advisor", "Read-only", "research, memory/files read, legal review, documents", "Issue spotting, licensing review, policy research; human legal review remains required."),
    ("Scribe", "Document Engineering", "Approval required", "memory/files read/write, document generation", "Technical documentation, SOPs, specifications, reports, release records."),
    ("Pixel", "Image and Visual Generation", "Approval required", "memory/files read/write, image generation", "Approved interface concepts, diagrams, product imagery, visual assets."),
    ("Blueprint", "Architecture and Blueprint Generation", "Read-only", "research, memory/files read, document/blueprint generation", "Architecture, data flow, workflow, and implementation blueprints."),
    ("Nexus", "API and Integration Engineering", "Approval required", "memory/files read/write, code/test execute", "Provider-neutral APIs, contracts, connectors, integration boundaries."),
    ("DataCore", "Data Engineering", "Approval required", "memory/files read/write, code execute, database admin", "Schemas, migrations, semantic-memory pipelines, retention, data operations."),
    ("Interface", "UX and Frontend Engineering", "Approval required", "research, memory/files read/write, code/browser/test", "Accessible UX and approved frontend implementation."),
    ("Release", "DevOps and Reliability", "Approval required", "diagnostics, memory/files read/write, code, deployment", "Build, deployment, observability, recovery, release processes."),
    ("Echo", "Voice and Experience Coordinator", "Read-only", "speech, avatar, memory/files read", "Local voice sessions and approved avatar experiences."),
]

TOOLS = [
    ("diagnostics", "Observability", "Low", "Read service health and runtime signals."),
    ("research", "Knowledge", "Low", "Investigate approved local or explicitly enabled sources."),
    ("investigation", "Development", "Low", "Correlate platform, task, audit, and workspace evidence."),
    ("memory_read", "Knowledge", "Low", "Retrieve approved workspace memory/context."),
    ("files_read", "Code", "Low", "Read files within the active workspace boundary."),
    ("files_write", "Code", "Medium", "Create or replace files in an authorized workspace."),
    ("code_execute", "Development", "High", "Run approved code in the isolated worker."),
    ("browser", "Development", "Medium", "Inspect approved browser surfaces when enabled."),
    ("speech", "Automation", "Low", "Use configured local STT and TTS."),
    ("avatar", "Automation", "Low", "Control local avatar presentation/state."),
    ("avatar_generate", "Automation", "Medium", "Generate local avatar artifacts from approved images."),
    ("security_scan", "Security", "Low", "Security inspection and dependency review."),
    ("compliance_review", "Compliance", "Low", "Map implementations to approved controls."),
    ("legal_review", "Compliance", "Low", "Legal and license issue spotting for human review."),
    ("test_execute", "Testing", "Medium", "Run authorized unit, integration, and regression tests."),
    ("document_generate", "Documentation", "Medium", "Create approved documents and reports."),
    ("image_generate", "Documentation", "Medium", "Create approved visual assets with configured local tooling."),
    ("blueprint_generate", "Documentation", "Low", "Produce architecture and workflow blueprints."),
    ("deployment", "Cloud", "Critical", "Governed release/deployment; separately authorized."),
    ("database_admin", "Database", "Critical", "Governed schema and administrative data operations."),
]

SERVICES = [
    ("app", "Core", "8080", "FastAPI, UI, REST, WebSocket, policy and orchestration"),
    ("worker", "Core", "internal 8092", "Allow-listed implementation actions in scoped workspaces"),
    ("postgres", "Core", "internal 5432", "Durable records plus pgvector schema"),
    ("redis", "Core", "internal 6379", "Priority queue, cache, transient state, kill messages"),
    ("ollama", "Core", "11434", "Default local open-weight model runtime"),
    ("portal", "Optional UI", "8082", "Legacy/holographic Gradio portal"),
    ("research-worker + SearXNG", "web-search", "internal", "Approved, single-purpose internet research route"),
    ("minio", "integrations", "9000/9001", "Optional S3-compatible artifact storage"),
    ("avatar3d", "avatar-local", "internal 8090", "Local TripoSR plus Blender generation"),
    ("speech", "speaking-avatar", "internal 8091", "Local Whisper-compatible STT and Kokoro/Piper TTS"),
    ("headtts", "speaking-avatar", "8882", "Optional TalkingHead/Kokoro presentation service"),
    ("echomimic", "avatar-lab", "7861", "GPU-oriented EchoMimic V2 laboratory"),
    ("openavatarchat", "live-avatar", "8282", "OpenAvatarChat + LiteAvatar experimental live runtime"),
]

API_GROUPS = [
    ("Health/config", "GET /api/health/live; /ready; /config", "Service diagnosis and UI configuration."),
    ("Agents", "GET/POST /api/agents; PATCH/DELETE /api/agents/{id}", "List, add, change permitted tools, or remove eligible agents."),
    ("Tasks", "POST/GET /api/tasks; POST /api/tasks/{id}/cancel", "Queue, observe, and cancel governed work."),
    ("Plans", "GET/POST /api/plans; POST /api/plans/{id}/decision", "Create and approve/reject implementation plans."),
    ("Lifecycle", "GET/POST /api/lifecycles; POST /api/lifecycles/{id}/transition", "Track test, sandbox, and production gates."),
    ("Worker", "GET /api/worker/health; POST /api/worker/actions", "Preview/write files and run allow-listed commands."),
    ("Workflows", "GET /api/workflows; POST /requests; POST /api/workflows", "Inspect, request, and register governed definitions."),
    ("Approvals", "POST /api/approvals; decision; external approvals", "Exact, expiring, one-time passcode decisions."),
    ("Research", "POST /api/research/search", "Consume approved search authorization through isolated broker."),
    ("Library", "GET /api/tool-library; changes; tool access requests", "Inspect registry and request controlled changes/access."),
    ("Knowledge", "GET /api/sources; /content; POST /requests", "View approved internal sources or request provenance review."),
    ("Workspace", "GET /api/workspace/tree; /file", "Read-only contained explorer and code preview."),
    ("Operations", "GET /api/security/posture; /audit; /metrics", "Security and operational evidence."),
    ("Control", "POST /api/control/kill-switch", "Stop/release agent execution."),
    ("Artifacts", "POST /api/artifacts", "Validated local uploads and context extraction."),
    ("Speech", "POST /api/speech/transcribe; /synthesize", "Local microphone transcription and voice output."),
    ("Avatar", "POST/GET/DELETE /api/avatar-generations", "Optional local avatar generation and governed deletion."),
    ("Progress", "WebSocket /api/ws", "Connected, task progress, avatar, and kill-switch events."),
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for idx, width in enumerate(widths):
        grid_cols[idx].set(qn("w:w"), str(round(width * 1440)))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    return run


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("ATLAS STUDIO  /  COMPLETE USER GUIDE"), 8, MUTED, True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Atlas Studio  |  Local-first operations  |  "), 8, MUTED)
    add_page_field(p)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        set_run(p.add_run(item))


def add_callout(doc, title, text, kind="info"):
    fill = {"info": PALE_CYAN, "warning": PALE_GOLD, "risk": PALE_RED, "success": "E8F5EF"}[kind]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    set_run(p.add_run(f"{title.upper()}  "), 9.5, NAVY, True)
    set_run(p.add_run(text), 9.5, INK)


def add_table(doc, headers, rows, widths, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        set_cell_shading(header.cells[idx], NAVY)
        p = header.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), font_size, WHITE, True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[idx], "F6F8FB")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run(p.add_run(str(value)), font_size, INK)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("A"), 52, CYAN, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("ATLAS STUDIO"), 30, NAVY, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Complete User Guide"), 19, BLUE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Admin  /  Operations  /  Developer"), 12, MUTED, True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Advanced Tooling, Lifecycle Automation, and Security"), 11.5, PURPLE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Build smarter. Operate safer. Scale confidently."), 11, MUTED, italic=True)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(f"Standalone Community Edition  |  Guide date: {date.today().isoformat()}"), 9, MUTED)
    doc.add_page_break()


def font(path, size):
    try:
        return ImageFont.truetype(path, size)
    except OSError:
        return ImageFont.load_default()


REGULAR = "C:/Windows/Fonts/arial.ttf"
BOLD = "C:/Windows/Fonts/arialbd.ttf"


def canvas(title, subtitle="", w=1800, h=980):
    image = Image.new("RGB", (w, h), "#F7FAFD")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((20, 20, w-20, h-20), 28, fill="#FFFFFF", outline="#C8D5E5", width=3)
    draw.text((70, 55), title, font=font(BOLD, 42), fill="#0A1830")
    if subtitle:
        draw.text((70, 112), subtitle, font=font(REGULAR, 22), fill="#61738B")
    return image, draw


def box(draw, xy, title, detail="", fill="#E8F8FB", outline="#20B9D5", title_color="#0A1830"):
    draw.rounded_rectangle(xy, 18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    tw = draw.textbbox((0,0), title, font=font(BOLD, 25))[2]
    draw.text(((x1+x2-tw)/2, y1+20), title, font=font(BOLD, 25), fill=title_color)
    if detail:
        lines = detail.split("\n")
        for i, line in enumerate(lines):
            sw = draw.textbbox((0,0), line, font=font(REGULAR, 18))[2]
            draw.text(((x1+x2-sw)/2, y1+62+i*24), line, font=font(REGULAR, 18), fill="#43556D")


def arrow(draw, start, end, color="#61738B", width=5):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    if abs(end[0]-start[0]) >= abs(end[1]-start[1]):
        s = 1 if end[0] > start[0] else -1
        draw.polygon([(x, y), (x-18*s, y-10), (x-18*s, y+10)], fill=color)
    else:
        s = 1 if end[1] > start[1] else -1
        draw.polygon([(x, y), (x-10, y-18*s), (x+10, y-18*s)], fill=color)


def make_diagrams():
    paths = {}
    img, d = canvas("Atlas Studio layered architecture", "Solid paths are core Community dependencies; optional profiles are explicitly enabled.")
    layers = [
        ("Experience", "Dashboard / Workspace / Voice / Workers", "#E8F8FB", "#20B9D5"),
        ("API and orchestration", "FastAPI / REST / WebSocket / LangGraph", "#EEF0FF", "#6957D9"),
        ("Policy and agents", "Security gates / Atlas / Forge / specialists", "#FFF4D9", "#C18A20"),
        ("Execution and intelligence", "Worker sandbox / Ollama / optional speech and avatar", "#E8F5EF", "#16845B"),
        ("Data and operations", "PostgreSQL + pgvector / Redis / artifacts / audit", "#F0F3F8", "#61738B"),
    ]
    y=170
    for title, detail, fill, outline in layers:
        box(d,(180,y,1620,y+115),title,detail,fill,outline); y+=145
    paths["architecture"] = ASSET_DIR / "architecture.png"; img.save(paths["architecture"])

    img, d = canvas("Site map", "Two navigation systems expose the same functional pages: categorized top navigation and persistent sidebar shortcuts.", h=1120)
    box(d,(690,150,1110,245),"Overview","Atlas assistant + movable status cards","#E8F8FB","#20B9D5")
    cats=[
        ("Build", ["Workspace Explorer","Code","Plans","Tasks","Implementation"]),
        ("Intelligence", ["Agents","Workflows","Tool Library","Knowledge","Sources of Truth"]),
        ("Assurance", ["Security","QA","Activity","Metrics"]),
        ("Experience", ["Workers","Avatar Lab"]),
        ("Settings", ["Environments","Profile foundation","Users foundation","OAuth foundation","Research approvals"]),
    ]
    xs=[50,400,750,1100,1450]
    for x,(name,items) in zip(xs,cats):
        box(d,(x,370,x+300,465),name,"\n".join(items[:1]),"#EEF0FF","#6957D9")
        arrow(d,(900,245),(x+150,370),"#9AAAC0",3)
        yy=500
        for item in items:
            d.rounded_rectangle((x+12,yy,x+288,yy+62),12,fill="#F7FAFD",outline="#C8D5E5",width=2)
            d.text((x+28,yy+19),item,font=font(REGULAR,18),fill="#18263A"); yy+=72
    paths["sitemap"] = ASSET_DIR / "sitemap.png"; img.save(paths["sitemap"])

    img, d = canvas("Governed implementation workflow", "Atlas coordinates; Forge changes files only after an exact user-approved plan and action.")
    nodes=[("User request","#E8F8FB","#20B9D5"),("Atlas plan","#EEF0FF","#6957D9"),("Passcode approval","#FFF4D9","#C18A20"),("Plan workspace","#F0F3F8","#61738B"),("Forge action","#E8F5EF","#16845B"),("Evidence","#E8F8FB","#20B9D5")]
    x=60
    for i,(name,fill,outline) in enumerate(nodes):
        box(d,(x,340,x+240,480),name,"",fill,outline)
        if i<len(nodes)-1: arrow(d,(x+240,410),(x+285,410))
        x+=285
    d.text((75,590),"Protected conditions",font=font(BOLD,24),fill="#0A1830")
    checks=["Atlas cannot receive mutating tools","Approval is exact, expiring, and one-time","Worker path remains inside the selected plan workspace","Network is denied by default","Diff, output, exit code, and hashes become evidence"]
    for i,item in enumerate(checks): d.text((100,640+i*48),f"✓  {item}",font=font(REGULAR,21),fill="#43556D")
    paths["implementation"] = ASSET_DIR / "implementation.png"; img.save(paths["implementation"])

    img, d = canvas("Development lifecycle", "Promotion depends on evidence, not an agent's assertion.")
    stages=[("Development","Approved plan\nScoped workspace"),("Test","Unit / integration\nRegression evidence"),("Sandbox","Isolated validation\nSecurity evidence"),("Production","Exact approval\nRelease record")]
    x=90
    colors=[("#E8F8FB","#20B9D5"),("#EEF0FF","#6957D9"),("#FFF4D9","#C18A20"),("#E8F5EF","#16845B")]
    for i,(name,detail) in enumerate(stages):
        box(d,(x,300,x+330,480),name,detail,*colors[i])
        if i<3: arrow(d,(x+330,390),(x+420,390),"#61738B",6)
        x+=420
    for i,label in enumerate(["Test gate","Sandbox gate","Production approval"]):
        d.text((405+i*420,530),label,font=font(BOLD,19),fill="#A23B3B")
    paths["lifecycle"] = ASSET_DIR / "lifecycle.png"; img.save(paths["lifecycle"])

    img, d = canvas("Research and development delivery", "External research is isolated from Forge's network-denied implementation workspace.", h=1100)
    top=["Intake","Scope","Research plan","Egress approval","Primary-source research","Option analysis"]
    bottom=["Decision record","Sandbox evaluation","Security + legal","QA validation","Isolated prototype","User approval"]
    for i,name in enumerate(top):
        x=55+i*285; box(d,(x,235,x+235,350),name,"", "#E8F8FB" if i<3 else "#FFF4D9", "#20B9D5" if i<3 else "#C18A20")
        if i<5: arrow(d,(x+235,292),(x+280,292),"#61738B",4)
    arrow(d,(1715,350),(1715,590),"#61738B",4)
    for i,name in enumerate(bottom):
        x=55+i*285; actual=bottom[::-1][i]
        box(d,(x,590,x+235,705),actual,"", "#E8F5EF" if i<3 else "#EEF0FF", "#16845B" if i<3 else "#6957D9")
        if i<5: arrow(d,(x+235,647),(x+280,647),"#61738B",4)
    d.text((90,820),"Agents: Atlas → Sage → Blueprint → user → Forge → Quanta → Sentinel + Counsel → user decision",font=font(BOLD,24),fill="#0A1830")
    d.text((90,875),"Outputs: research brief, source register, architecture options, prototype, test evidence, risk decision",font=font(REGULAR,21),fill="#43556D")
    paths["rnd"] = ASSET_DIR / "rnd.png"; img.save(paths["rnd"])

    img, d = canvas("Data flow and ownership", "Each store has a distinct purpose and workspace boundary.")
    box(d,(80,330,360,480),"Browser","REST requests\nWebSocket events","#E8F8FB","#20B9D5")
    box(d,(520,330,840,480),"FastAPI control plane","Validation / policy\nLangGraph / audit","#EEF0FF","#6957D9")
    boxes=[("PostgreSQL + pgvector","Durable records\nSemantic-memory schema"),("Redis","Queue / cache\nkill-switch messages"),("Artifacts","Validated local files\nor optional MinIO"),("Ollama","Local model\nresponse stream")]
    ys=[170,340,510,680]
    for (title,detail),y in zip(boxes,ys):
        box(d,(1120,y,1660,y+125),title,detail,"#F0F3F8","#61738B")
        arrow(d,(840,405),(1120,y+62),"#61738B",4)
    arrow(d,(360,405),(520,405),"#61738B",6)
    paths["data"] = ASSET_DIR / "data.png"; img.save(paths["data"])

    img, d = canvas("Security trust boundary", "Model output never grants permissions; deterministic API policy and one-time approvals do.")
    box(d,(70,310,340,450),"User","Intent + approval","#E8F8FB","#20B9D5")
    box(d,(455,310,735,450),"API policy","Tool allow-list\nRisk + scope checks","#EEF0FF","#6957D9")
    box(d,(850,310,1130,450),"Agent","Named identity\nPrompted behavior","#FFF4D9","#C18A20")
    box(d,(1245,310,1535,450),"Worker sandbox","no-new-privileges\nlimits + scoped path","#E8F5EF","#16845B")
    for x1,x2 in [(340,455),(735,850),(1130,1245)]: arrow(d,(x1,380),(x2,380),"#61738B",5)
    d.rounded_rectangle((1200,590,1690,780),18,fill="#FCEBEA",outline="#A23B3B",width=3)
    d.text((1250,630),"Denied by default",font=font(BOLD,25),fill="#A23B3B")
    for i,t in enumerate(["Internet route","Host filesystem escape","Container socket","Unapproved production access"]): d.text((1260,678+i*28),t,font=font(REGULAR,18),fill="#43556D")
    d.rounded_rectangle((110,590,1050,805),18,fill="#F7FAFD",outline="#C8D5E5",width=3)
    d.text((160,630),"Enforced controls",font=font(BOLD,25),fill="#0A1830")
    controls=["Atlas read-only boundary","Passcode approval","Workspace containment","Upload validation","Audit events","Kill switch"]
    for i,t in enumerate(controls):
        x=170+(i%3)*285; y=690+(i//3)*58; d.text((x,y),f"✓ {t}",font=font(REGULAR,18),fill="#43556D")
    paths["security"] = ASSET_DIR / "security.png"; img.save(paths["security"])

    img, d = canvas("Agent capability distribution", "The chart counts registered capabilities assigned to each built-in agent.", h=1100)
    counts=[]
    for name,_,_,tool_text,_ in AGENTS:
        counts.append((name,len([x for x in tool_text.split(",") if x.strip()])))
    max_count=max(v for _,v in counts)
    y=175
    for name,count in counts:
        d.text((80,y),name,font=font(BOLD,19),fill="#18263A")
        d.rounded_rectangle((300,y,300+count/max_count*1250,y+25),12,fill="#20B9D5" if name=="Atlas" else "#6957D9")
        d.text((1580,y),str(count),font=font(BOLD,18),fill="#61738B"); y+=54
    paths["agent_chart"] = ASSET_DIR / "agent_chart.png"; img.save(paths["agent_chart"])
    return paths


def add_figure(doc, path, caption, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split(".", 1)[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    set_run(p.add_run(caption), 8.5, MUTED, italic=True)


def chapter(doc, number, title, intro=None):
    if len(doc.paragraphs) > 8:
        doc.add_page_break()
    doc.add_heading(f"{number}. {title}", level=1)
    if intro:
        add_body(doc, intro)


def build():
    diagrams = make_diagrams()
    doc = Document()
    configure_document(doc)
    def add_heading(text, level=2):
        return doc.add_heading(text, level=level)
    title_page(doc)

    doc.add_heading("Guide scope and status legend", level=1)
    add_body(doc, "This guide documents the Atlas Studio repository as implemented on the guide date. It is written for the local platform owner acting in Admin, Operations, or Developer capacity. It does not treat interface placeholders or optional Compose profiles as core production capabilities.")
    add_table(doc, ["Label", "Meaning"], [
        ("AVAILABLE", "Implemented in the current application or core Compose stack."),
        ("OPTIONAL PROFILE", "Implemented but disabled until its Compose profile and configuration are enabled."),
        ("FOUNDATION / PLANNED", "Visible schema, UI, or design foundation; not an active end-to-end capability."),
        ("EXPERIMENTAL", "Source-integrated feature with hardware, model, or maturity constraints."),
    ], [1.65,4.85], 9)
    add_callout(doc,"Identity boundary","Admin, Ops, and Developer are personas in this release. Community mode currently uses one local platform-owner session. Multi-user accounts, enforced RBAC, and Google OAuth are foundations/planned, not active controls.","warning")
    add_callout(doc,"Decision integrity","Agents are instructed not to invent missing facts, files, results, permissions, requirements, or preferences. When required information is missing, the agent should identify the gap, ask a direct question, and wait. Deterministic authorization policy still governs actions independently of model output.","info")

    doc.add_heading("Contents", level=1)
    contents = [
        "1. Platform at a glance", "2. Role-based operating model", "3. Site map and navigation", "4. Overview dashboard", "5. Working with Atlas", "6. Agents", "7. Skills and Tool Library", "8. Workflows", "9. Plans and approvals", "10. Workspace and Code", "11. Tasks and priority", "12. Implementation worker", "13. QA, Sandbox, and Environments", "14. Security", "15. Operations and Metrics", "16. Knowledge and Sources of Truth", "17. Artifacts and uploads", "18. Voice, Workers, and avatar features", "19. Administration and configuration", "20. Role-based use cases", "21. Admin how-tos", "22. Operations how-tos", "23. Developer how-tos", "24. Architecture and data flows", "25. API reference", "26. Libraries and service inventory", "27. Troubleshooting", "28. Known boundaries and glossary"
    ]
    add_bullets(doc, contents)

    chapter(doc,1,"Platform at a glance","Atlas Studio is a standalone, local-first multi-agent engineering control plane. Its design separates orchestration, implementation, security, persistence, and optional experience services.")
    add_figure(doc,diagrams["architecture"],"Figure 1. Implemented layered architecture.")
    add_table(doc,["Mode","What runs","Credentials"],[
        ("Community (default)","App, worker, PostgreSQL/pgvector, Redis, Ollama, local artifacts","No external key required"),
        ("Integrations","Community core plus explicitly enabled adapters such as MinIO","Only the enabled adapter's local settings"),
        ("Optional profiles","Web search, speech/avatar, avatar generation, avatar lab, live avatar","No paid API required; models may download separately"),
    ],[1.45,3.65,1.4],8.7)
    add_bullets(doc,["Primary UI: http://localhost:8080","API documentation: http://localhost:8080/api/docs","Optional Gradio portal: http://localhost:8082","Ollama API: http://localhost:11434","OpenAvatarChat experimental UI: http://localhost:8282 when enabled"])

    chapter(doc,2,"Role-based operating model")
    add_table(doc,["Persona","Primary responsibility","Typical pages","Authorization expectation"],[
        ("Admin","Configuration, agent/tool governance, sources, workflow registry, approvals","Settings, Agents, Tool Library, Workflows, Sources, Security","Owns local passcode and protected decisions"),
        ("Operations","Availability, queue, audit, metrics, incidents, kill switch","Overview, Tasks, Activity, Metrics, Security, QA","Can stop work; should not authorize code without owner mandate"),
        ("Developer","Inspect code, propose plans, use Forge, test and promote evidence","Workspace, Code, Plans, Implementation, QA, Sandbox","Requires approved plan/workspace and exact action approval"),
    ],[1.0,2.25,1.75,1.5],8.2)
    add_callout(doc,"Not RBAC","These personas are operational guidance. The server does not yet authenticate separate Admin, Ops, and Developer identities.","risk")
    add_heading("Suggested separation of duties",level=2)
    add_bullets(doc,["The Admin approves plans, restricted capability changes, internet research, avatar deletion, and production transitions.","Operations observes health and evidence, manages incidents, and uses the kill switch when necessary.","Developers prepare scoped requests and implementation evidence but do not self-approve protected actions.","Sentinel, Quanta, Verity, and Counsel provide independent review; their output does not replace qualified human decisions."])

    chapter(doc,3,"Site map and navigation")
    add_figure(doc,diagrams["sitemap"],"Figure 2. Current navigation and page hierarchy.")
    add_table(doc,["Top category","Pages","Primary audience"],[
        ("Overview","Assistant feed, Task Orchestrator, Active Tasks, Atlas Status","All"),
        ("Build","Workspace Explorer, Code, Plans, Tasks, Implementation","Developer, Admin"),
        ("Intelligence","Agents, Workflows, Tool Library, Knowledge, Sources of Truth","Admin, Developer"),
        ("Assurance","Security, QA, Activity, Metrics","Operations, Admin"),
        ("Experience","Workers, Avatar Lab","Admin, experience developer"),
        ("Settings","Environments, profile/user foundations, research approvals, runtime config","Admin"),
    ],[1.15,3.7,1.65],8.8)
    add_body(doc,"The left sidebar provides persistent shortcuts to Dashboard, Workspace, Agents, Tools, Knowledge, Analytics, and Settings plus a live built-in-agent roster. The top navigation exposes the full category map. The profile menu links to profile/account, users/access, and Settings foundations.")

    chapter(doc,4,"Overview dashboard")
    add_body(doc,"Overview embeds the Atlas Assistant and three operational cards: Task Orchestrator, Active Tasks, and Atlas Status.")
    add_heading("Reorganize Overview cards",level=2)
    add_bullets(doc,["Drag a card by its header handle to a new vertical position.","For keyboard control, focus a card header and press Ctrl (or Command) plus an arrow key.","The order is stored in that browser's local storage.","Select Reset card order to restore the default. The central assistant panel remains fixed."] , numbered=True)
    add_heading("Status interpretation",level=2)
    add_table(doc,["Indicator","Meaning","Response"],[
        ("Healthy / Online","Core dependencies and model respond","Proceed"),
        ("Degraded","UI/API live but at least one dependency unavailable","Open Metrics or Activity; inspect readiness"),
        ("Model unavailable","Ollama reachable but configured model absent or timed out","Pull/choose model; verify RAM and logs"),
        ("Kill switch engaged","New tasks blocked; active/queued work cancelled","Investigate, then release only when safe"),
    ],[1.45,2.9,2.15],8.8)

    chapter(doc,5,"Working with Atlas")
    add_body(doc,"Atlas is the read-only orchestrator in the chain User → Atlas → Forge → specialist agents. Atlas may diagnose, research, investigate, read approved memory, and read workspace files. It cannot write files, execute code, run tests, deploy, or administer databases.")
    add_heading("Ask a useful question",level=2)
    add_bullets(doc,["State the outcome: what decision, diagnosis, plan, or artifact you need.","Name the workspace, component, file, environment, or task when known.","State constraints: no external access, performance target, deadline, or applicable source.","Request evidence: files inspected, logs, tests, risks, and unresolved questions.","If a required fact is missing, answer Atlas's clarification before authorizing action."] , numbered=True)
    add_callout(doc,"Example","Inspect the authentication flow in the current workspace. Identify why tool creation returns 403, cite the exact files and evidence, propose a plan, and ask before assuming any missing requirement.","info")
    add_heading("Text, files, and voice",level=2)
    add_bullets(doc,["Use the composer for text instructions and the attachment controls for supported local files.","Uploads are validated and stored locally; attachment context is extracted where supported.","When the speaking-avatar profile is active, the microphone transcribes locally and Atlas can synthesize a cleaned natural-language reply.","Speech output removes code-like symbols and non-speech material where possible; the message feed remains the record."])

    chapter(doc,6,"Agents")
    add_figure(doc,diagrams["agent_chart"],"Figure 3. Approximate built-in capability assignment counts from the current agent registry.")
    add_table(doc,["Agent","Role","Boundary","Capabilities","Purpose"],AGENTS,[0.72,1.25,1.0,1.65,1.88],7.1)
    add_heading("Add a named agent",level=2)
    add_bullets(doc,["Open Intelligence → Agents and select Add new agent.","Enter a unique name, role, and responsibility/boundary description.","Select only the skills/tools required for its purpose.","Keep Read-only selected unless implementation is necessary.","Keep Require my authorization selected. Agents with mutating tools are rejected if authorization is disabled.","Approve the protected creation action when prompted, then verify the card and audit event."],numbered=True)
    add_heading("Modify or remove an agent",level=2)
    add_body(doc,"Tool changes use the same governed API boundary. Atlas cannot be given mutating tools. System agents may be protected from deletion; only delete an eligible custom agent after confirming it is not needed for active tasks or workflow ownership.")

    chapter(doc,7,"Skills and Tool Library")
    add_callout(doc,"Terminology","In the current UI, “skills and tools” are the registered ToolId capabilities below. Atlas Studio does not yet provide a general package installer or execute arbitrary external skill code. Manual workflows may reference an existing skill or external resource as metadata; those references are not downloaded or executed automatically.","warning")
    add_table(doc,["Capability","Category","Risk","Function"],TOOLS,[1.25,1.05,0.65,3.55],8.0)
    add_heading("Use the Tool Library",level=2)
    add_bullets(doc,["Search or filter the platform registry by category.","Open a capability to inspect provider, source, risk, data access, permissions, environments, assigned agents, and audit requirement.","Select Request / add to record a governed access request for an agent and environment.","Treat deployment and database_admin as restricted production capabilities requiring separate authorization.","A request does not automatically change an agent's permissions or install software."],numbered=True)
    add_heading("Request a library change",level=2)
    add_body(doc,"The control plane supports add, update, and remove library-change records. Record the tool identifier, name, description, reason, and action; then complete technical, security, provenance, and licensing review before any implementation.")

    chapter(doc,8,"Workflows")
    add_table(doc,["Workflow","Owner","Status","Nodes / purpose"],[
        ("Governed agent task","Atlas","Active","Policy → model → deterministic completion"),
        ("Software feature delivery","Forge","Designed","Plan → approve → implement → QA → security → release"),
        ("Security remediation","Sentinel","Designed","Verify → approve → fix → test → retest → release"),
        ("Research to decision","Sage","Designed","Frame → research → compare → review → decide"),
        ("R&D delivery","Sage","Designed","Sourced research through isolated prototype, validation, and decision"),
    ],[1.6,0.8,0.75,3.35],8.3)
    add_figure(doc,diagrams["rnd"],"Figure 4. Built-in Research and Development delivery workflow.")
    add_heading("Request a workflow",level=2)
    add_bullets(doc,["Open Intelligence → Workflows.","Under Request a workflow, enter a name, outcome/evidence goal, proposed owner, and optional local skill/library/URL references.","Submit the request. This creates intake metadata and grants no execution authority.","Review the proposed nodes, sources, tool needs, approvals, outputs, failure paths, and evidence before implementation."],numbered=True)
    add_heading("Register a manual workflow",level=2)
    add_bullets(doc,["Choose a stable lowercase identifier, name, owner, source type, description, and ordered node list.","Provide the reviewed reference; external URLs remain untrusted metadata.","Approve the exact workflow_definition action with the local passcode.","The definition is registered pending security review; registration does not install or execute external code."],numbered=True)

    chapter(doc,9,"Plans and approvals")
    add_figure(doc,diagrams["implementation"],"Figure 5. User-controlled implementation path.")
    add_heading("Create and approve a plan",level=2)
    add_bullets(doc,["Open Build → Plans and enter a title, implementation agent, priority, and requested outcome.","Review the generated plan steps and scope. Do not approve vague paths, commands, or acceptance criteria.","Approve with the local passcode only when the exact plan is acceptable.","Atlas Studio creates a plan-specific workspace copy. Wait until its status is ready.","Use that workspace for implementation tasks and worker actions. Rejecting a plan grants no authority."],numbered=True)
    add_heading("Approval properties",level=2)
    add_bullets(doc,["Exact: bound to action, target, and payload.","Expiring: unusable after its configured lifetime.","One-time: consumed by the approved action.","Audited: request, decision, and consumption are recorded.","Passcode is submitted for the decision and is not stored in the browser."])
    add_callout(doc,"Default credential","Change ATLAS_STUDIO_APPROVAL_PASSCODE before using write, execution, internet, avatar deletion, or production approvals.","risk")

    chapter(doc,10,"Workspace Explorer and Code")
    add_body(doc,"The app mounts the repository at /workspace read-only. The explorer hides secrets and generated dependency folders, rejects traversal and symlink escapes, and limits previews by configuration.")
    add_heading("Inspect a file",level=2)
    add_bullets(doc,["Open Build → Workspace Explorer.","Expand directories within the mounted project.","Select a supported file. Atlas Studio audits the read and opens it in Code.","Use line numbers and breadcrumb details for review. The Code surface is read-only.","To edit, create/approve a plan and use Forge in Implementation rather than editing through Code."],numbered=True)
    add_callout(doc,"Secrets","The explorer is a convenience boundary, not a substitute for secret hygiene. Keep credentials outside source control and avoid network exposure until authentication/TLS is implemented.","warning")

    chapter(doc,11,"Tasks and priority")
    add_body(doc,"Tasks use a durable Redis priority queue and PostgreSQL persistence. Status flows through queued, running, completed, failed, or cancelled. WebSocket events update the UI.")
    add_table(doc,["Priority","Use","Behavior"],[
        ("Critical","Active incident or urgent security containment","Highest queue position; still requires all approvals"),
        ("High","Implementation and important engineering work","Forge normal/low requests are elevated to high"),
        ("Normal","Routine read-only analysis or planned work","Default"),
        ("Low","Non-urgent background analysis","Lowest queue position"),
    ],[0.9,3.3,2.3],8.8)
    add_heading("Add or cancel a task",level=2)
    add_bullets(doc,["Open Build → Tasks and enter title, agent, priority, and instructions.","Read-only agents may accept direct diagnostic/research tasks.","Implementation-capable agents require an approved plan and ready plan workspace; direct bypass attempts are rejected.","Watch status in Tasks, Overview, Activity, or Metrics.","Cancel queued/running work from the task controls; the kill switch cancels all active/queued work."],numbered=True)

    chapter(doc,12,"Implementation worker")
    add_body(doc,"The separate worker receives narrow actions over an internal network. It has a read-only container root, no new privileges, no Linux capabilities, 512 MB memory, 1 CPU, 128 PID limit, and a plan-scoped workspace. It does not receive a generic Docker socket or internet route.")
    add_table(doc,["Action","Input","Evidence returned"],[
        ("Preview file change","Path + complete proposed content","Diff and file hashes; no write"),
        ("Write approved file","Exact path + complete content + approval","Write result, diff/hashes"),
        ("Run approved code","Allow-listed command arguments + approval","Output and exit code"),
        ("Run approved tests","Allow-listed test arguments + approval","Test output, exit code, evidence"),
    ],[1.45,2.55,2.5],8.8)
    add_heading("Run a safe Forge action",level=2)
    add_bullets(doc,["Confirm worker health is available.","Select Forge or another authorized implementation agent and a ready plan workspace.","Choose Preview file change first; enter a relative workspace path and complete replacement content.","Review the proposed diff and hashes.","Select the intended write/execute/test action, check the exact authorization box, and continue.","Review the passcode modal summary; approve only if action, path/command, and workspace match.","Retain worker output as lifecycle evidence."],numbered=True)

    chapter(doc,13,"QA, Sandbox, and Environments")
    add_figure(doc,diagrams["lifecycle"],"Figure 6. Evidence-gated development lifecycle.")
    add_table(doc,["Environment","Purpose","Default boundary","Gate"],[
        ("Development","Plan, inspect, implement in scoped workspace","Workspace-specific permissions","Approved plan/workspace"),
        ("Test","Unit, integration, regression, security, performance checks","Authorized test execution","Completed-task evidence"),
        ("Sandbox","Isolated behavior and security validation","Network none; limits; dropped capabilities","Test or security evidence"),
        ("Production","Release/deployment state","Separately governed","Exact one-time production approval"),
    ],[1.05,2.2,1.85,1.4],8.4)
    add_heading("Promote lifecycle evidence",level=2)
    add_bullets(doc,["Create a lifecycle for an approved plan.","Attach completed implementation/test evidence before requesting Test completion.","Review Quanta results and unresolved failures.","Move to Sandbox only with qualifying test or security evidence.","Have Sentinel review relevant security risk and Counsel/Verity review applicable legal/compliance issues.","Request production transition only after acceptance criteria and rollback/operational evidence are complete.","Approve the exact production transition; record the release decision."],numbered=True)

    chapter(doc,14,"Security")
    add_figure(doc,diagrams["security"],"Figure 7. Security control and trust boundary.")
    add_table(doc,["Layer","Current state","Control"],[
        ("Identity/session","Planned","Authentication, session expiry, CSRF before non-loopback exposure"),
        ("Agent permissions","Enforced","Server-side tool allow-lists; protected Atlas boundary"),
        ("Human authorization","Enforced","Protected actions require exact approval"),
        ("Workspace","Enforced","Resolved-path containment; read-only app mount"),
        ("Sandbox","Enforced","Network none, limits, capabilities dropped"),
        ("Uploads","Enforced","Filename, extension/type, content, and size validation"),
        ("Audit/evidence","Enforced","Tasks, agents, approvals, controls, uploads, lifecycle events"),
        ("Secrets/integrations","Enforced by mode","Optional integrations disabled; telemetry off in Community"),
    ],[1.25,1.05,4.2],8.4)
    add_heading("Use the kill switch",level=2)
    add_bullets(doc,["Select Stop all agents when execution must halt.","The API blocks new tasks, removes queued work, cancels queued/running tasks, publishes a kill message, writes an audit event, and updates clients.","Investigate service, model, worker, and audit evidence.","Release the switch only after the unsafe or failed condition is understood."],numbered=True)
    add_callout(doc,"Network exposure","Before exposing Atlas Studio beyond localhost, add local authentication/RBAC, CSRF protection, a TLS reverse proxy, secret rotation, backup/restore, and network policy. These are not complete today.","risk")

    chapter(doc,15,"Operations and Metrics")
    add_table(doc,["Metric panel","What it shows","Operator question"],[
        ("Task performance","Status distribution, queue depth, durations, success rate","Is work flowing or failing?"),
        ("Service health","API, model, worker, Postgres, Redis, speech/avatar state","Which dependency is degraded?"),
        ("Model runtime","Provider, model, readiness, timeout, token ceiling","Is response latency/model availability acceptable?"),
        ("Agent governance","System/custom/read-only/implementation counts","Are capabilities appropriately constrained?"),
        ("Security posture","Atlas boundary, sandbox, telemetry, audit outcomes","Are controls enforced?"),
        ("Runtime/storage","Process memory/load, artifacts, database/cache state","Is capacity healthy?"),
        ("Recent tasks/audit","Latest activity and accountability records","What changed and who/what acted?"),
        ("Tool coverage","Capability assignments","Are tools over- or under-assigned?"),
    ],[1.3,2.65,2.55],8.4)
    add_heading("Daily operations checklist",level=2)
    add_bullets(doc,["Confirm /api/health/live and /api/health/ready.","Review degraded services and model readiness.","Review queue depth, running task age, failures, and cancellations.","Check kill-switch state and recent denied/failed audit events.","Confirm artifact storage growth and persistent volume health.","Review pending approvals and expiring research permissions.","Validate backups before high-risk changes (backup automation remains an operator responsibility)."])

    chapter(doc,16,"Knowledge and Sources of Truth")
    add_body(doc,"Knowledge search currently filters the approved local source catalog, not a full semantic-memory search service. The registered authoritative sources are README.md, SECURITY.md, and IMPLEMENTATION.md when available in the workspace.")
    add_table(doc,["Precedence (highest first)","Source type"],[(1,"Applicable law and regulation"),(2,"Official government sources"),(3,"Regulators and standards organizations"),(4,"Organization policies and procedures"),(5,"Approved internal documentation"),(6,"Applicable compliance requirements"),(7,"Official vendor documentation"),(8,"Other verified sources"),(9,"General web and unverified information")],[1.5,5.0],9)
    add_heading("Request a new source",level=2)
    add_bullets(doc,["Open Sources of Truth and select Request source addition.","Enter name, authority, source type, location, and optional jurisdiction/version.","Submit for provenance review.","The request is pending and is not authoritative until reviewed for identity, scope, currency, jurisdiction, integrity, conflicts, and approval.","Atlas should flag unresolved source conflicts instead of silently choosing one."],numbered=True)
    add_callout(doc,"Semantic memory","PostgreSQL includes a workspace-scoped pgvector memory schema with 1,024-dimensional vectors. End-to-end embedding ingestion and user-facing semantic retrieval remain a hardening/incremental implementation item.","warning")

    chapter(doc,17,"Artifacts and uploads")
    add_body(doc,"Artifacts use the local filesystem by default. MinIO is optional in Integrations mode. Uploads are reduced to safe basenames, checked against an allow-list, limited by size, resolved inside the artifact root, and audited.")
    add_table(doc,["Category","Accepted examples"],[
        ("Documents/data","txt, md, json, csv, pdf, rtf, odt, docx, xlsx, pptx"),
        ("Images/3D/audio","png, jpg/jpeg, webp, glb/gltf, wav, mp3"),
        ("Code/config","py, js, ts/tsx/jsx, html, css, yaml/yml, toml, sql, sh, ps1"),
        ("Default limit","25 MB general uploads; speech and avatar endpoints have separate limits"),
    ],[1.4,5.1],9)
    add_bullets(doc,["Do not upload secrets, private keys, or unlicensed personal imagery.","Confirm permission before avatar image processing.","An uploaded file is context, not executable authority.","Executable extensions and path traversal names are rejected."])

    chapter(doc,18,"Voice, Workers, and avatar features")
    add_table(doc,["Feature","Status","Use / constraint"],[
        ("Workers by delos","Available UI","Local GLB/WebGL worker identities; Atlas image and source references included"),
        ("Local speech","Optional profile","Whisper-compatible STT; Kokoro/Piper American female TTS defaults"),
        ("Image-to-3D","Optional profile","TripoSR + Blender; front image drives geometry; not a production human rig"),
        ("Speaking avatar","Optional profile","TalkingHead/Three.js foundation; requires rig/morph-compatible character"),
        ("EchoMimic V2","Experimental GPU lab","CUDA/NVIDIA and official model weights required"),
        ("LiteAvatar/OpenAvatarChat","Experimental","Large first-start downloads; CPU path may be slow; hardware dependent"),
    ],[1.35,1.2,3.95],8.3)
    add_heading("Local voice how-to",level=2)
    add_bullets(doc,["Start the speaking-avatar profile and wait for speech health.","Open the Atlas chat surface and grant microphone permission.","Press the microphone, speak, and stop/submit the turn as shown.","Verify the transcript before relying on it for an implementation decision.","Atlas text appears in the feed; synthesized audio plays through the browser when local TTS succeeds.","If audio fails, inspect speech service health and the X-Atlas-Voice-Backend response header."],numbered=True)
    add_callout(doc,"Lifelike boundary","The repository includes experimental/open-source avatar paths, but it does not currently deliver a turnkey photorealistic, full-body, low-latency neural avatar on CPU-only hardware.","warning")

    chapter(doc,19,"Administration and configuration")
    add_table(doc,["Configuration area","Default","Admin action"],[
        ("Mode","community","Use integrations only for explicitly enabled local adapters"),
        ("Model","Ollama / qwen3:8b","Pull model separately; select a smaller manifest when hardware constrained"),
        ("Artifacts","filesystem / 25 MB","Enable MinIO only with profile and changed credentials"),
        ("Sandbox","docker; network none; 512m; 1 CPU; 128 PIDs","Keep deny-by-default; tune only with evidence"),
        ("Telemetry","false","Keep disabled unless an explicit local policy permits it"),
        ("Approval passcode","development default","Change before protected actions or exposure"),
        ("OAuth","disabled","Foundation only; do not treat as active sign-in"),
    ],[1.45,2.05,3.0],8.4)
    add_heading("Start and verify the core",level=2)
    add_bullets(doc,["Copy .env.example to .env and change local passwords/passcode.","Run docker compose up -d --build.","Pull the configured Ollama model separately.","Run docker compose ps until core services are healthy.","Open http://localhost:8080 and inspect Metrics.","Call /api/health/live and /api/health/ready for precise component state."],numbered=True)

    chapter(doc,20,"Role-based use cases")
    add_table(doc,["ID","Persona","Goal","Primary path","Outcome"],[
        ("UC-A1","Admin","Create a least-privilege custom agent","Agents → add → choose tools → approve","Named audited agent"),
        ("UC-A2","Admin","Register a reviewed workflow","Workflows → manual definition → passcode","Pending-security-review definition"),
        ("UC-A3","Admin","Approve external research","Settings → query/purpose/domains → passcode","One-time 15-minute route"),
        ("UC-A4","Admin","Add a source candidate","Sources → request → provenance review","Pending, not authoritative"),
        ("UC-O1","Ops","Diagnose degradation","Metrics → readiness → logs","Identified dependency/failure"),
        ("UC-O2","Ops","Stop unsafe execution","Stop all agents → audit → investigate","Tasks cancelled; execution locked"),
        ("UC-O3","Ops","Review task performance","Metrics/Activity/Tasks","Queue and failure evidence"),
        ("UC-D1","Developer","Diagnose a code problem","Atlas → Workspace/Code → plan","Evidence-backed plan"),
        ("UC-D2","Developer","Implement a feature","Plan → approve → Forge preview/write → test","Scoped change + evidence"),
        ("UC-D3","Developer","Promote a release","QA → Sandbox → Production approval","Traceable lifecycle decision"),
        ("UC-D4","Developer","Evaluate technology","R&D workflow → approved research → prototype","Sourced recommendation"),
    ],[0.55,0.8,1.35,2.2,1.6],7.8)

    chapter(doc,21,"Admin how-tos")
    procedures = [
        ("Change an agent's tools",["Open Agents and select the named agent.","Compare requested capability with the Tool Library risk/data-access details.","Remove unnecessary tools; never assign mutating tools to Atlas.","Approve protected changes when required.","Verify the refreshed card and audit event."]),
        ("Approve internet research",["Start the web-search profile if the local broker is required.","In Settings, enter the exact query, business purpose, and optional allowed domains.","Review the passcode modal and approve.","Run the search before the 15-minute expiry.","Confirm the approval was consumed and retain sources in the research record."]),
        ("Govern a source",["Request the source with authority, type, location, jurisdiction, and version.","Validate ownership, currency, integrity, applicability, and conflicts.","Do not treat the submitted item as authoritative until the review process is implemented/completed.","Record the decision and update dependent workflows/documentation."]),
        ("Safely enable an optional profile",["Review license, model, hardware, port, storage, and privacy requirements.","Set only the documented environment values.","Build/start that profile without changing Community dependencies.","Wait for its health check and review logs.","Disable it if degraded; core Atlas should continue operating."]),
    ]
    for name,steps in procedures:
        doc.add_heading(name,level=2); add_bullets(doc,steps,numbered=True)

    chapter(doc,22,"Operations how-tos")
    procedures = [
        ("Investigate a slow or timed-out response",["Check Metrics → Model runtime and Service health.","Confirm Ollama health and that the configured model is present.","Review task duration, timeout, model token ceiling, host RAM/CPU, and competing optional services.","Stop resource-heavy experimental profiles if core work is affected.","Use a smaller local model only after confirming quality needs."]),
        ("Recover from degraded core services",["Use readiness to identify Postgres, Redis, Ollama, or worker failure.","Inspect the specific service logs and container state.","Restore the dependency; avoid deleting persistent volumes during diagnosis.","Confirm readiness, queue recovery, and task/audit consistency.","Document incident cause and corrective action."]),
        ("Handle an unsafe or runaway task",["Engage Stop all agents.","Verify queued/running tasks become cancelled and new task creation is locked.","Inspect audit, task, worker, and workspace evidence.","Correct permissions, plan scope, command allow-list, or service condition.","Release the switch and use a low-risk diagnostic task before resuming."]),
        ("Monitor storage",["Review Metrics → Runtime & storage.","Inspect artifact and model volume growth.","Back up Postgres and required artifacts/models according to local policy.","Use governed deletion for generated avatar artifacts; do not manually remove unknown volume data."])
    ]
    for name,steps in procedures:
        doc.add_heading(name,level=2); add_bullets(doc,steps,numbered=True)

    chapter(doc,23,"Developer how-tos")
    procedures = [
        ("Diagnose before changing code",["Ask Atlas for evidence and exact file references.","Browse the mounted workspace and open relevant files in Code.","Request clarification for missing requirements or environment facts.","Create a plan with acceptance criteria, test scope, risk, and rollback expectations.","Do not jump directly to Forge for mutating work."]),
        ("Implement a change with Forge",["Approve the plan and wait for a ready plan workspace.","Use Preview file change with a relative path and complete content.","Review diff and hashes.","Approve and execute the exact file write.","Run targeted tests through the worker.","Attach output and exit code to lifecycle evidence."]),
        ("Run R&D",["Frame a decision question, constraints, evaluation criteria, and required sources.","Let Sage prepare the research plan; approve egress only if needed.","Have Blueprint compare architectural options.","Approve only a bounded prototype.","Forge builds it without internet; Quanta validates; Sentinel and Counsel review.","Record the user-owned adopt/reject/iterate decision."]),
        ("Add a workflow from a library or skill",["Review the local library/skill documentation and license outside the execution path.","Create a workflow request with desired outcome and reference.","Translate it into named nodes, owners, inputs, outputs, gates, failure states, and evidence.","Register the definition with passcode approval.","Complete security review before activation; external code is not auto-installed."])
    ]
    for name,steps in procedures:
        doc.add_heading(name,level=2); add_bullets(doc,steps,numbered=True)

    chapter(doc,24,"Architecture and data flows")
    add_figure(doc,diagrams["data"],"Figure 8. Runtime data flow and store ownership.")
    add_table(doc,["Store","Durability","Data"],[
        ("PostgreSQL","Durable","Workspaces, agents, tasks, plans, lifecycles, workflow runs/steps/approvals/events, audit, artifact metadata, memory schema"),
        ("pgvector","Durable extension","Workspace-scoped 1,024-dimensional semantic-memory vectors"),
        ("Redis","Durable queue + transient","Priority queue, task snapshots, cache, coordination, kill messages"),
        ("Filesystem","Durable volume","Default uploaded/generated artifacts"),
        ("MinIO","Optional durable","S3-compatible artifacts in Integrations profile"),
        ("Plan workspace volume","Durable until archived","Isolated implementation copies bound to approved plans"),
        ("Ollama volume","Durable","Separately downloaded local model weights"),
    ],[1.45,1.25,3.8],8.3)
    add_heading("Task event flow",level=2)
    add_bullets(doc,["Browser submits REST request.","API validates agent, kill switch, plan/workspace, tools, risk, and authorization.","Task persists to PostgreSQL and priority is queued in Redis.","LangGraph policy node runs before the model node.","Ollama streams local output; the API emits WebSocket deltas/progress.","Terminal state and audit evidence persist; the UI reconciles task status."],numbered=True)

    chapter(doc,25,"API reference")
    add_table(doc,["Area","Endpoints","Use"],API_GROUPS,[1.1,3.45,1.95],7.8)
    add_callout(doc,"API security","The API is designed for localhost Community use. Do not treat the current local-owner session as sufficient authentication for LAN or internet exposure.","risk")

    chapter(doc,26,"Libraries and service inventory")
    add_heading("Application libraries",level=2)
    add_table(doc,["Library","Role"],[
        ("FastAPI / Uvicorn","REST, WebSocket, validation hosting"),("Pydantic Settings","Validated environment configuration"),("HTTPX","Local provider/service calls"),("asyncpg","PostgreSQL persistence"),("redis-py","Queue/cache/control messages"),("LangGraph + PostgreSQL checkpointer","Policy-first agent workflow and durable checkpoints"),("python-multipart","Uploads"),("Gradio","Optional portal"),("pypdf / python-docx / openpyxl / python-pptx / Pillow","Local document and image context handling")
    ],[2.4,4.1],8.8)
    add_heading("Container services",level=2)
    add_table(doc,["Service","Profile","Port","Purpose"],SERVICES,[1.25,1.1,0.8,3.35],7.8)
    add_heading("Model/provider boundary",level=2)
    add_body(doc,"Ollama is the default native provider. An OpenAI-compatible protocol adapter supports self-hosted llama.cpp, vLLM, and compatible Transformers servers. No vendor SDK is required. Model manifests reference suggested Qwen general/coder models and TripoSR; weights are not bundled.")

    chapter(doc,27,"Troubleshooting")
    add_table(doc,["Symptom","Likely cause","Action"],[
        ("localhost:8080 refused","App container stopped or port unavailable","docker compose ps; inspect app logs; restart core"),
        ("UI unchanged after rebuild","Old image/container or browser cache","Rebuild/recreate app; Ctrl+F5"),
        ("Ollama 404","Configured model not pulled","ollama list; pull exact configured model"),
        ("Model timeout","Model too large, host contention, or timeout","Inspect resources; stop optional profiles; use smaller model or tune timeout"),
        ("Task 403/409","Missing authorization, plan, or ready workspace","Approve plan; select its ready workspace; retry exact action"),
        ("Worker unavailable","worker service unhealthy/token mismatch","Check worker health/logs and shared token"),
        ("Speech 503","Speaking profile disabled/unhealthy","Start speech profile; inspect /health and app TTS/STT URLs"),
        ("No browser audio","Autoplay/permission or invalid TTS response","Interact with page; inspect status/header; test synthesized WAV"),
        ("Avatar service unhealthy","Model bootstrap/hardware/build failure","Review profile logs and requirements; disable without affecting core"),
        ("OpenAvatarChat starting/unhealthy","Large model initialization or resource pressure","Follow logs; wait for downloads; confirm RAM/GPU constraints"),
        ("Docker API 500","Docker Desktop engine/config issue","Restart Docker Desktop; verify engine; retry pull/compose"),
        ("Card reorder not visible","Old static bundle/cache","Rebuild app and hard refresh; drag card header handle"),
    ],[1.4,2.4,2.7],7.9)
    add_heading("Diagnostic commands",level=2)
    for command in ["docker compose ps","docker compose logs --tail 100 app","docker compose logs --tail 100 worker","docker compose exec ollama ollama list","Invoke-RestMethod http://localhost:8080/api/health/live","Invoke-RestMethod http://localhost:8080/api/health/ready"]:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); set_run(p.add_run(command),9,NAVY,font="Consolas")

    chapter(doc,28,"Known boundaries and glossary")
    add_heading("Known implementation boundaries",level=2)
    add_bullets(doc,["Community mode is a single local-owner workspace; multi-user authentication and enforced RBAC are not active.","Google OAuth is disabled and foundational only.","Knowledge search filters registered local sources; complete semantic-memory ingestion/retrieval is not exposed end to end.","Manual workflow registration stores definitions; only the governed-agent-task graph is active universally. Designed workflows require further execution-node implementation.","Generic Docker socket execution is intentionally absent; a narrow Docker action broker remains future work.","Production deployment and database administration are restricted capabilities, not a general self-service production control plane.","Avatar services are optional/experimental and do not guarantee photorealistic, low-latency full-body behavior on current hardware.","Local-first does not automatically mean encrypted, authenticated, backed up, or safe for network exposure; those controls must be implemented and operated explicitly."])
    add_heading("Glossary",level=2)
    add_table(doc,["Term","Meaning"],[
        ("Approval","Exact, time-limited, single-use user decision for a protected action."),("Artifact","Validated uploaded or generated local file."),("Community mode","No-required-key, local core deployment."),("Evidence","Recorded diff, hash, test result, exit code, source, audit event, or decision supporting a gate."),("Forge","Primary implementation agent; requires user authorization."),("LangGraph","OSS workflow engine used for policy-first orchestration."),("Plan workspace","Isolated workspace copy bound to one approved plan."),("Source of Truth","Reviewed authority with provenance, version, jurisdiction, and applicability."),("Tool / skill","Registered capability that can be assigned to a named agent."),("Workspace","Mounted project boundary visible read-only to the app/Atlas."),
    ],[1.5,5.0],8.8)
    add_callout(doc,"Operating principle","Build smarter. Operate safer. Scale confidently. Keep Atlas read-only, keep implementation scoped to Forge and approved agents, demand evidence at every gate, and ask the user when a decision lacks required information.","success")

    doc.core_properties.title = "Atlas Studio Complete User Guide"
    doc.core_properties.subject = "Admin, Operations, and Developer guide"
    doc.core_properties.author = "Atlas Studio"
    doc.core_properties.keywords = "Atlas Studio, user guide, admin, operations, developer, security, workflows"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
